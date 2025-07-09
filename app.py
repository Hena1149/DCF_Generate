import streamlit as st
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx import Document
from openai import AzureOpenAI
import io
import time

# Configuration de la page
st.set_page_config(
    page_title="Générateur de DCF",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS personnalisé qui respecte les couleurs du config.toml
st.markdown("""
<style>
    :root {
        --primary-color: #4CAF50;
        --secondary-color: #e3f2fd;
        --background-color: #f8f9fa;
        --text-color: #2c3e50;
        --accent-color: #2196F3;
        --border-color: #ced4da;
    }
    
    .stApp {
        background-color: var(--background-color);
        color: var(--text-color);
        font-family: 'sans serif';
    }
    
    .stTextInput>div>div>input, 
    .stTextArea>div>div>textarea,
    .stSelectbox>div>div>select {
        background-color: white !important;
        color: var(--text-color) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 6px !important;
        padding: 8px 12px !important;
    }
    
    .stTextInput>div>div>input:focus, 
    .stTextArea>div>div>textarea:focus {
        border-color: var(--primary-color) !important;
        box-shadow: 0 0 0 2px rgba(76, 175, 80, 0.2) !important;
    }
    
    .stButton>button {
        background-color: var(--primary-color) !important;
        color: white !important;
        border: none !important;
        border-radius: 6px !important;
        padding: 10px 24px !important;
        font-weight: 500 !important;
        transition: all 0.3s ease !important;
    }
    
    .stButton>button:hover {
        background-color: #3e8e41 !important;
        transform: translateY(-1px);
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    }
    
    .stFileUploader>div>div>button {
        background-color: var(--primary-color) !important;
        color: white !important;
        border-radius: 6px !important;
    }
    
    .sidebar .sidebar-content {
        background-color: var(--secondary-color) !important;
        border-right: 1px solid var(--border-color);
    }
    
    h1, h2, h3, h4, h5, h6 {
        color: var(--text-color) !important;
    }
    
    .stProgress>div>div>div>div {
        background-color: var(--primary-color) !important;
    }
    
    .stAlert {
        background-color: var(--secondary-color) !important;
        border-left: 4px solid var(--primary-color);
    }
    
    .stMarkdown {
        color: var(--text-color) !important;
    }
    
    .download-btn {
        background-color: var(--accent-color) !important;
        margin-top: 10px !important;
    }
    
    .download-btn:hover {
        background-color: #0d8bf2 !important;
    }
    
    .header-container {
        background: linear-gradient(135deg, var(--primary-color) 0%, #3e8e41 100%);
        padding: 2rem;
        border-radius: 0 0 10px 10px;
        margin-bottom: 2rem;
    }
    
    .header-title {
        color: white !important;
        margin: 0;
    }
    
    .header-subtitle {
        color: rgba(255,255,255,0.9) !important;
        margin: 0.5rem 0 0;
    }
    
    .info-box {
        background-color: var(--secondary-color);
        border-left: 4px solid var(--accent-color);
        padding: 1rem;
        border-radius: 6px;
        margin-bottom: 1rem;
    }
    
    .success-box {
        background-color: rgba(76, 175, 80, 0.1);
        border-left: 4px solid var(--primary-color);
        padding: 1rem;
        border-radius: 6px;
        margin-bottom: 1rem;
    }
    
    .error-box {
        background-color: rgba(244, 67, 54, 0.1);
        border-left: 4px solid #f44336;
        padding: 1rem;
        border-radius: 6px;
        margin-bottom: 1rem;
    }
    
    .stExpander {
        border: 1px solid var(--border-color) !important;
        border-radius: 6px !important;
    }
    
    .stExpander .stExpanderHeader {
        background-color: var(--secondary-color) !important;
    }
    
    [data-testid="stHeader"] {
        background-color: rgba(255,255,255,0) !important;
    }
    
    [data-testid="stToolbar"] {
        display: none !important;
    }
</style>
""", unsafe_allow_html=True)

def read_file(uploaded_file):
    """Lit le contenu d'un fichier uploadé (PDF, TXT ou DOCX)"""
    try:
        if uploaded_file.type == "application/pdf":
            text = ""
            with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
                for page in doc:
                    text += page.get_text()
            return text
        
        elif uploaded_file.type == "text/plain":
            return uploaded_file.read().decode("utf-8")
        
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = DocxDocument(io.BytesIO(uploaded_file.read()))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        else:
            st.error("Format non supporté. Utilisez un fichier .pdf, .txt ou .docx.")
            return None
    except Exception as e:
        st.error(f"Erreur lors de la lecture du fichier: {str(e)}")
        return None

# def generate_prompt(cdc_text):
#     """Génère le prompt pour GPT à partir du texte du CDC"""
#     return f"""
# Tu es un assistant expert en conception fonctionnelle de systèmes d'information, et tu dois rédiger un Dossier de Conception Fonctionnelle (DCF) à partir d'un cahier des charges (CDC) fourni ci-dessous.

# Le DCF que tu vas rédiger doit **respecter rigoureusement la structure suivante**, issue du guide d'élaboration DDI M IT 02.02 :

# ---

# ### 1. CADRE GENERAL
# 1.1. Présentation générale du système (objectifs, fonctions globales)
# 1.2. Références (documents applicables et références)
# 1.3. Environnement (positionnement dans le SI, environnement technique)
# 1.4. Terminologie et sigles utilisés

# ### 2. ARCHITECTURE FONCTIONNELLE
# 2.1. Modules fonctionnels (découpage, description des modules)
# 2.2. Synoptique fonctionnel (flux entre fonctions)

# ### 3. SPECIFICATIONS FONCTIONNELLES
# Pour chaque module identifié :
# - Nom du module
# - Pour chaque fonction :
#   - Définition (objectif de la fonction)
#   - Identification (code, acteur, déclencheur, conséquences IHM et traitement)
#   - Description du processus :
#     - Entrées
#     - Traitement
#     - Sorties
#     - Règles de gestion (Pas d'Abréviation écrit la règle de gestion)

# ### 4. REPRISE DE L'EXISTANT
# 4.1. Procédure de reprise
# 4.2. Contraintes de reprise

# ### 5. RECAPITULATIF DES REGLES DE GESTION
# Tableau récapitulatif avec fonction associée à chaque règle.

# ### 6. VISA DE VALIDATION
# Présentation des aspects validés et les parties prenantes concernées.

# ---

# Tu dois **extraire, analyser et structurer le contenu du CDC suivant** pour produire automatiquement un DCF de qualité conforme à cette structure, en tenant compte :
# - des besoins exprimés,
# - des règles de gestion métier,
# - des exigences fonctionnelles,
# - des contraintes techniques,
# - des modules évoqués.

# Voici le contenu du CDC :

# \"\"\"{cdc_text[:15000]}\"\"\"

# Rédige maintenant un DCF complet et bien formaté à partir de ce CDC.
# """

def generate_prompt(cdc_text):
    """Génère le prompt pour GPT à partir du texte du CDC"""
    return f"""
Tu es un assistant expert en conception fonctionnelle de systèmes d'information, et tu dois rédiger un Dossier de Conception Fonctionnelle (DCF) détaillé et complet à partir d'un cahier des charges (CDC) fourni ci-dessous.

Le DCF que tu vas rédiger doit **respecter rigoureusement la structure suivante**, issue du guide d'élaboration DDI M IT 02.02, en fournissant des informations précises et exhaustives pour chaque section :

---

### 1. CADRE GENERAL
1.1. Présentation générale du système
   - Objectifs stratégiques et opérationnels
   - Périmètre fonctionnel précis
   - Finalité du système
   - Bénéfices attendus
   - Publics cibles

1.2. Références
   - Documents normatifs (liste complète)
   - Standards applicables
   - Contraintes réglementaires
   - Références aux documents projets

1.3. Environnement
   - Architecture technique détaillée
   - Systèmes connectés (interfaces)
   - Contraintes d'intégration
   - Prérequis matériels/logiciels
   - Environnement de déploiement

1.4. Terminologie et sigles
   - Glossaire complet avec définitions
   - Liste des acronymes avec explications
   - Termes techniques spécifiques

### 2. ARCHITECTURE FONCTIONNELLE
2.1. Modules fonctionnels
   - Découpage modulaire détaillé
   - Responsabilités de chaque module
   - Interactions entre modules
   - Spécificités techniques

2.2. Synoptique fonctionnel
   - Diagramme textuel des flux
   - Séquencement des opérations
   - Points d'intégration critiques
   - Flux principaux et secondaires

### 3. SPECIFICATIONS FONCTIONNELLES (À DÉTAILLER POUR CHAQUE MODULE)
Pour chaque module identifié :
- Nom du module et version
- Description approfondie :
  * Finalité et portée
  * Contraintes spécifiques
  * Hypothèses techniques

Pour chaque fonction :
  - Définition complète :
    * Objectif métier
    * Valeur ajoutée
    * Critères de succès

  - Identification précise :
    * Code unique (norme de nommage)
    * Acteurs concernés (rôles)
    * Déclencheurs (événements)
    * Préconditions et postconditions
    * Impacts IHM détaillés

  - Description du processus :
    * Entrées : format, source, validation
    * Traitement : algorithme, logique métier
    * Sorties : format, destination, qualité
    * Règles de gestion : formulation complète sans abréviation
    * Cas d'erreur et gestion des exceptions
    * Contrôles de qualité

### 4. REPRISE DE L'EXISTANT
4.1. Procédure de reprise
   - Stratégie de migration
   - Plan de conversion
   - Nettoyage des données
   - Validation post-migration

4.2. Contraintes de reprise
   - Compatibilités
   - Anomalies connues
   - Limitations techniques
   - Périmètre exclu

### 5. RECAPITULATIF DES REGLES DE GESTION
Tableau structuré contenant :
- Identifiant unique de la règle
- Libellé complet et non ambigu
- Module/fonction associée
- Source métier
- Critère d'application
- Exemples concrets
- Exceptions éventuelles

### 6. VISA DE VALIDATION
- Liste des validations requises
- Responsables par domaine
- Critères d'acceptation
- Preuves de validation
- Planning de recette

---

**Directives spécifiques :**
1. Analyse minutieusement le CDC pour extraire toutes les exigences implicites et explicites
2. Structure le contenu de manière logique et progressive
3. Utilise un langage technique précis mais accessible
4. Fournis des exemples concrets quand nécessaire
5. Identifie clairement les dépendances entre composants
6. Mentionne les contraintes et limitations de manière transparente
7. Propose des recommandations pour les aspects critiques

**Approche rédactionnelle :**
- Style professionnel et normatif
- Phrases complètes et structurées
- Terminologie cohérente
- Numérotation précise des éléments
- Mise en forme claire avec des paragraphes aérés

Voici le contenu du CDC à analyser :

\"\"\"{cdc_text[:30000]}\"\"\"

Génère maintenant un DCF exhaustif, en développant particulièrement :
- Les règles de gestion avec leur logique complète
- Les scénarios d'utilisation typiques
- Les cas limites à prendre en compte
- Les interfaces système détaillées
- Les contraintes de performance
"""

def call_gpt(prompt, api_key, endpoint, deployment):
    """Appelle l'API Azure OpenAI pour générer le DCF"""
    try:
        client = AzureOpenAI(
            api_key=api_key,
            api_version="2024-02-15-preview",
            azure_endpoint=endpoint
        )

        response = client.chat.completions.create(
            model=deployment,
            messages=[
                {"role": "system", "content": "Tu es un expert en conception de systèmes logiciels."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Erreur lors de l'appel à l'API OpenAI: {str(e)}")
        return None

def save_dcf_to_word(text, filename="DCF_Généré.docx"):
    """Sauvegarde le DCF dans un fichier Word"""
    doc = Document()
    for line in text.split('\n'):
        if line.strip() == "":
            continue
        if line.strip().startswith('#') or line.strip().startswith("1.") or line.strip().startswith("2.") or line.strip().startswith("3."):
            doc.add_heading(line.strip(), level=1)
        else:
            doc.add_paragraph(line.strip())
    
    # Sauvegarde en mémoire
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def save_dcf_to_txt(text, filename="DCF_Généré.txt"):
    """Sauvegarde le DCF dans un fichier texte"""
    buffer = io.BytesIO()
    buffer.write(text.encode('utf-8'))
    buffer.seek(0)
    return buffer

def main():
    """Fonction principale de l'application Streamlit."""
    # Header avec dégradé de couleur
    st.markdown("""
    <div class="header-container">
        <h1 class="header-title">Générateur de Dossier de Conception Fonctionnelle (DCF)</h1>
        <p class="header-subtitle">Transformez votre cahier des charges en un DCF professionnel en quelques clics</p>
    </div>
    """, unsafe_allow_html=True)
    
    with st.expander("Instructions d'utilisation", expanded=True):
        st.markdown("""
        <div class="info-box">
            <h4 style="margin-top: 0;">Comment utiliser cette application :</h4>
            <ol>
                <li><strong>Configurez votre API Azure OpenAI</strong> dans la barre latérale</li>
                <li><strong>Téléversez votre cahier des charges</strong> (PDF, TXT ou DOCX)</li>
                <li><strong>Générez le DCF</strong> en cliquant sur le bouton</li>
                <li><strong>Téléchargez le résultat</strong> au format Word ou texte</li>
            </ol>
        </div>
        """, unsafe_allow_html=True)
    
    with st.sidebar:
        st.markdown("""
        <div style="border-bottom: 2px solid var(--primary-color); padding-bottom: 10px; margin-bottom: 20px;">
            <h2 style="color: var(--text-color); margin: 0;">⚙️ Configuration</h2>
        </div>
        """, unsafe_allow_html=True)
        
        st.subheader("API Azure OpenAI")
        api_key = st.text_input("Clé API", type="password", help="La clé API pour accéder au service Azure OpenAI")
        endpoint = st.text_input("Endpoint", value="https://chat-genai.openai.azure.com/", help="L'URL du endpoint Azure OpenAI")
        deployment = st.text_input("Modèle", value="gpt-4o", help="Le nom du modèle déployé dans Azure OpenAI")
        
        st.markdown("---")
        
        st.subheader("Options")
        show_prompt = st.checkbox("Afficher le prompt envoyé à l'API", value=False)
        show_raw_output = st.checkbox("Afficher la sortie brute de l'API", value=False)
    
    # Zone de téléchargement du fichier avec style amélioré
    st.subheader("Téléversement du fichier")
    uploaded_file = st.file_uploader(
        " ",
        type=["pdf", "txt", "docx"],
        help="Format acceptés: PDF, TXT ou DOCX",
        label_visibility="collapsed"
    )
    
    # Bouton de génération avec icône
    generate_button = st.button("Générer le DCF", type="primary", use_container_width=True)
    
    if generate_button:
        if not uploaded_file:
            st.markdown("""
            <div class="error-box">
                <p style="margin: 0;">Veuillez téléverser un fichier CDC.</p>
            </div>
            """, unsafe_allow_html=True)
            return
        
        if not api_key or not endpoint or not deployment:
            st.markdown("""
            <div class="error-box">
                <p style="margin: 0;">Veuillez configurer les paramètres Azure OpenAI dans la barre latérale.</p>
            </div>
            """, unsafe_allow_html=True)
            return
        
        try:
            with st.spinner("Lecture du fichier en cours..."):
                cdc_text = read_file(uploaded_file)
                time.sleep(1)
            
            if not cdc_text.strip():
                st.markdown("""
                <div class="error-box">
                    <p style="margin: 0;">Le fichier semble vide ou n'a pas pu être lu correctement.</p>
                </div>
                """, unsafe_allow_html=True)
                return
            
            with st.spinner("Génération du prompt..."):
                prompt = generate_prompt(cdc_text)
                time.sleep(1)
                if show_prompt:
                    with st.expander(" Prompt envoyé à l'API"):
                        st.code(prompt)
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for percent in range(0, 101, 10):
                status_text.text(f" Génération en cours... {percent}%")
                progress_bar.progress(percent)
                time.sleep(0.1)
            
            start_time = time.time()
            dcf_result = call_gpt(prompt, api_key, endpoint, deployment)
            elapsed_time = time.time() - start_time
            
            progress_bar.empty()
            status_text.empty()
            
            st.markdown(f"""
            <div class="success-box">
                <h4 style="margin-top: 0;"> DCF généré avec succès !</h4>
                <p>Temps de traitement : {elapsed_time:.2f} secondes</p>
            </div>
            """, unsafe_allow_html=True)
            
            if show_raw_output:
                with st.expander("Sortie brute de l'API"):
                    st.code(dcf_result)
            
            # Affichage du résultat avec onglets
            tab1, tab2 = st.tabs(["Aperçu du DCF", " Téléchargement"])
            
            with tab1:
                st.subheader("Résultat - Dossier de Conception Fonctionnelle")
                st.markdown(dcf_result)
            
            with tab2:
                st.subheader("Options de téléchargement")
                col1, col2 = st.columns(2)
                with col1:
                    word_buffer = save_dcf_to_word(dcf_result)
                    st.download_button(
                        label="Télécharger en Word",
                        data=word_buffer,
                        file_name="DCF_Généré.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                with col2:
                    txt_buffer = save_dcf_to_txt(dcf_result)
                    st.download_button(
                        label=" Télécharger en TXT",
                        data=txt_buffer,
                        file_name="DCF_Généré.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
            
        except Exception as e:
            st.markdown(f"""
            <div class="error-box">
                <h4 style="margin-top: 0;"> Une erreur est survenue</h4>
                <p>{str(e)}</p>
            </div>
            """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
